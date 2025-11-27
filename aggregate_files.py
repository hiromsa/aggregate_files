#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import argparse
import time
import re
from pathlib import Path
from urllib.parse import urlparse, urljoin
from typing import List, Set, Optional, Tuple
import threading
from datetime import datetime, timedelta
import multiprocessing
from concurrent.futures import ProcessPoolExecutor, as_completed

import requests
from bs4 import BeautifulSoup
import pdfplumber
import openpyxl
import docx


class FileAggregator:
    def __init__(self, input_source: str, output_file: str):
        self.input_source = input_source
        self.output_file = output_file
        self.visited_urls: Set[str] = set()
        self.base_domain = None
        self.base_path = None
        
        # 進捗表示用の変数
        self.total_files = 0
        self.processed_files = 0
        self.start_time = None
        self.current_file = ""
        self.lock = threading.Lock()
        
        # スキップするファイル/ディレクトリのリスト
        self.skip_patterns = [
            r'bin$', r'obj$', r'\.git$', r'\.vs$', r'__pycache__$',
            r'node_modules$', r'\.exe$', r'\.dll$', r'\.pdb$',
            r'\.zip$', r'\.tar\.gz$', r'\.log$', r'\.jpg$',
            r'\.jpeg$', r'\.png$', r'\.ico$', r'\.css$', r'\.js$'
        ]
        
        # 未対応のファイル形式
        self.unsupported_formats = {'.xls', '.doc'}
        
        # 並列処理用の設定
        self.max_workers = min(multiprocessing.cpu_count(), 8)  # 最大8ワーカーに制限
    
    def is_web_url(self, source: str) -> bool:
        """入力ソースがWeb URLかどうかを判定"""
        return source.startswith(('http://', 'https://'))
    
    def should_skip_file(self, file_path: str) -> bool:
        """ファイルをスキップすべきか判定"""
        for pattern in self.skip_patterns:
            if re.search(pattern, file_path, re.IGNORECASE):
                return True
        return False
    
    def is_unsupported_format(self, file_path: str) -> bool:
        """未対応のファイル形式か判定"""
        ext = Path(file_path).suffix.lower()
        return ext in self.unsupported_formats
    
    def extract_pdf_text(self, file_path: str) -> str:
        """PDFファイルからテキストを抽出"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except Exception as e:
            return f"[ERROR: Failed to extract PDF content: {str(e)}]"
    
    def extract_xlsx_text(self, file_path: str) -> str:
        """Excelファイルからテキストを抽出"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            return text
        except Exception as e:
            return f"[ERROR: Failed to extract Excel content: {str(e)}]"
    
    def extract_docx_text(self, file_path: str) -> str:
        """Wordファイルからテキストを抽出"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"[ERROR: Failed to extract Word content: {str(e)}]"
    
    def update_progress(self):
        """進捗状況を表示"""
        with self.lock:
            if self.total_files > 0:
                progress_percent = (self.processed_files / self.total_files) * 100
                elapsed = datetime.now() - self.start_time if self.start_time else timedelta(0)
                
                if self.processed_files > 0 and elapsed.total_seconds() > 0:
                    # 処理速度（ファイル/秒）
                    rate = self.processed_files / elapsed.total_seconds()
                    # 残りファイル数
                    remaining_files = self.total_files - self.processed_files
                    # 残り時間予測（秒）
                    remaining_seconds = remaining_files / rate if rate > 0 else 0
                    # 予測完了時刻
                    eta = datetime.now() + timedelta(seconds=remaining_seconds)
                    eta_str = eta.strftime("%Y-%m-%d %H:%M:%S")
                else:
                    eta_str = "計算中..."
                
                print(f"\r進捗: {progress_percent:.1f}% ({self.processed_files}/{self.total_files} ファイル) "
                      f"現在処理中: {self.current_file} "
                      f"完了予定時刻: {eta_str}", end="", flush=True)
    
    def process_local_file(self, file_path: str, relative_path: str) -> str:
        """ローカルファイルを処理"""
        with self.lock:
            self.current_file = relative_path
        
        if self.should_skip_file(relative_path):
            with self.lock:
                self.processed_files += 1
                self.update_progress()
            return ""
        
        if self.is_unsupported_format(relative_path):
            with self.lock:
                self.processed_files += 1
                self.update_progress()
            return f"# File: {relative_path}\n```text\n[WARNING: The file format ({Path(relative_path).suffix}) is not supported. Content was skipped.]\n```\n\n"
        
        ext = Path(relative_path).suffix.lower()
        
        try:
            if ext == '.pdf':
                content = self.extract_pdf_text(file_path)
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext == '.xlsx':
                content = self.extract_xlsx_text(file_path)
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext == '.docx':
                content = self.extract_docx_text(file_path)
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.csv', '.sql']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                result = f"# File: {relative_path}\n```{ext[1:] if ext else 'text'}\n{content}\n```\n\n"
            else:
                # テキストファイルとして読み込みを試みる
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
        except Exception as e:
            result = f"# File: {relative_path}\n```text\n[ERROR: Failed to read file: {str(e)}]\n```\n\n"
        
        with self.lock:
            self.processed_files += 1
            self.update_progress()
        
        return result
    
    @staticmethod
    def process_file_worker(args: Tuple[str, str, List[str], Set[str]]) -> Tuple[str, str]:
        """ワーカープロセスでファイルを処理する静的メソッド"""
        file_path, relative_path, skip_patterns, unsupported_formats = args
        
        # スキップパターンのチェック
        for pattern in skip_patterns:
            if re.search(pattern, relative_path, re.IGNORECASE):
                return relative_path, ""
        
        # 未対応フォーマットのチェック
        ext = Path(relative_path).suffix.lower()
        if ext in unsupported_formats:
            return relative_path, f"# File: {relative_path}\n```text\n[WARNING: The file format ({ext}) is not supported. Content was skipped.]\n```\n\n"
        
        # ファイル処理
        try:
            if ext == '.pdf':
                # PDF処理
                try:
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        content = text
                except Exception as e:
                    content = f"[ERROR: Failed to extract PDF content: {str(e)}]"
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext == '.xlsx':
                # Excel処理
                try:
                    workbook = openpyxl.load_workbook(file_path)
                    text = ""
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        text += f"Sheet: {sheet_name}\n"
                        for row in sheet.iter_rows(values_only=True):
                            row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                            if row_text.strip():
                                text += row_text + "\n"
                        text += "\n"
                    content = text
                except Exception as e:
                    content = f"[ERROR: Failed to extract Excel content: {str(e)}]"
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext == '.docx':
                # Word処理
                try:
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                    content = text
                except Exception as e:
                    content = f"[ERROR: Failed to extract Word content: {str(e)}]"
                result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
            elif ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.csv', '.sql']:
                # テキストファイル処理
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    result = f"# File: {relative_path}\n```{ext[1:] if ext else 'text'}\n{content}\n```\n\n"
                except Exception as e:
                    result = f"# File: {relative_path}\n```text\n[ERROR: Failed to read file: {str(e)}]\n```\n\n"
            else:
                # その他のファイル形式
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    result = f"# File: {relative_path}\n```text\n{content}\n```\n\n"
                except Exception as e:
                    result = f"# File: {relative_path}\n```text\n[ERROR: Failed to read file: {str(e)}]\n```\n\n"
        except Exception as e:
            result = f"# File: {relative_path}\n```text\n[ERROR: Failed to process file: {str(e)}]\n```\n\n"
        
        return relative_path, result
    
    def process_local_directory_parallel(self, root_path: str) -> str:
        """ローカルディレクトリを並列処理"""
        root_path = Path(root_path).resolve()
        
        # まず全ファイル数をカウント
        print("ファイル数をカウント中...")
        all_files = [(str(file_path), str(file_path.relative_to(root_path)))
                    for file_path in root_path.rglob('*')
                    if file_path.is_file() and not self.should_skip_file(str(file_path.relative_to(root_path)))]
        
        self.total_files = len(all_files)
        self.processed_files = 0
        self.start_time = datetime.now()
        
        print(f"並列処理開始: 合計 {self.total_files} ファイル（{self.max_workers} ワーカー）")
        
        # 結果を格納する辞書（順序保持用）
        results = {}
        
        # 並列処理の実行
        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            # ワーカーに渡す引数を準備
            tasks = [(file_path, relative_path, self.skip_patterns, self.unsupported_formats)
                    for file_path, relative_path in all_files]
            
            # タスクをサブミット
            future_to_path = {executor.submit(self.process_file_worker, task): task[1] for task in tasks}
            
            # 完了したタスクから結果を収集
            for future in as_completed(future_to_path):
                relative_path = future_to_path[future]
                try:
                    path, result = future.result()
                    results[path] = result
                    
                    # 進捗更新
                    with self.lock:
                        self.processed_files += 1
                        self.current_file = relative_path
                        self.update_progress()
                except Exception as e:
                    error_result = f"# File: {relative_path}\n```text\n[ERROR: Worker process failed: {str(e)}]\n```\n\n"
                    results[relative_path] = error_result
                    
                    # 進捗更新
                    with self.lock:
                        self.processed_files += 1
                        self.current_file = relative_path
                        self.update_progress()
        
        # 結果を元の順序で結合
        final_result = ""
        for _, relative_path in all_files:
            if relative_path in results:
                final_result += results[relative_path]
        
        print()  # 進捗表示の後に改行
        return final_result
    
    def process_local_directory(self, root_path: str) -> str:
        """ローカルディレクトリを再帰的に処理"""
        # 並列処理を使用
        return self.process_local_directory_parallel(root_path)
    
    def is_same_domain(self, url: str) -> bool:
        """URLが同じドメインか判定"""
        parsed = urlparse(url)
        return parsed.netloc == self.base_domain
    
    def is_under_base_path(self, url: str) -> bool:
        """URLがベースパス配下か判定"""
        parsed = urlparse(url)
        return parsed.path.startswith(self.base_path)
    
    def download_binary_file(self, url: str, session: requests.Session) -> Optional[str]:
        """バイナリファイルをダウンロードしてテキストを抽出"""
        try:
            response = session.get(url, stream=True)
            response.raise_for_status()
            
            # 一時ファイルに保存
            temp_file = f"temp_{os.path.basename(url)}"
            with open(temp_file, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            # ファイル拡張子に基づいてテキスト抽出
            ext = Path(url).suffix.lower()
            if ext == '.pdf':
                content = self.extract_pdf_text(temp_file)
            elif ext == '.xlsx':
                content = self.extract_xlsx_text(temp_file)
            elif ext == '.docx':
                content = self.extract_docx_text(temp_file)
            else:
                content = f"[WARNING: Binary file format ({ext}) is not supported for content extraction.]"
            
            # 一時ファイルを削除
            os.remove(temp_file)
            
            return content
        except Exception as e:
            return f"[ERROR: Failed to download or process {url}: {str(e)}]"
    
    def crawl_web_page(self, url: str, session: requests.Session) -> str:
        """Webページをクロールしてコンテンツを抽出"""
        if url in self.visited_urls:
            return ""
        
        with self.lock:
            self.current_file = url
            self.processed_files += 1
            self.update_progress()
        
        self.visited_urls.add(url)
        
        try:
            response = session.get(url)
            response.raise_for_status()
            
            # コンテンツタイプを確認
            content_type = response.headers.get('content-type', '').lower()
            
            # バイナリファイルの場合
            if any(ext in url.lower() for ext in ['.pdf', '.xlsx', '.docx']):
                content = self.download_binary_file(url, session)
                return f"# URL: {url}\n```text\n{content}\n```\n\n"
            
            # HTMLページの場合
            if 'text/html' in content_type:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # スクリプトとスタイルを除去
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # テキストコンテンツを抽出
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                result = f"# URL: {url}\n```text\n{text}\n```\n\n"
                
                # 同じドメイン内のリンクを探索
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    absolute_url = urljoin(url, href)
                    
                    if self.is_same_domain(absolute_url) and self.is_under_base_path(absolute_url):
                        if absolute_url not in self.visited_urls:
                            time.sleep(1)  # サーバー負荷軽減のための遅延
                            result += self.crawl_web_page(absolute_url, session)
                
                return result
            else:
                # その他のコンテンツタイプ
                return f"# URL: {url}\n```text\n[WARNING: Content type '{content_type}' is not supported for text extraction.]\n```\n\n"
                
        except Exception as e:
            return f"# URL: {url}\n```text\n[ERROR: Failed to crawl {url}: {str(e)}]\n```\n\n"
    
    def process_web_source(self, start_url: str) -> str:
        """Webソースを処理"""
        parsed = urlparse(start_url)
        self.base_domain = parsed.netloc
        self.base_path = '/'.join(parsed.path.split('/')[:-1]) if '/' in parsed.path else '/'
        
        # Webクロール用の初期化
        self.total_files = 50  # 見積もり値（実際のページ数はクロール中に変動）
        self.processed_files = 0
        self.start_time = datetime.now()
        
        print(f"Webクロール開始: {start_url}")
        
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        result = self.crawl_web_page(start_url, session)
        
        print()  # 進捗表示の後に改行
        return result
    
    def aggregate(self):
        """メインの集約処理"""
        print(f"開始: {self.input_source}")
        start_time = datetime.now()
        
        if self.is_web_url(self.input_source):
            content = self.process_web_source(self.input_source)
        else:
            content = self.process_local_directory(self.input_source)
        
        # 出力ファイルに書き込み
        print("出力ファイルを作成中...")
        with open(self.output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        end_time = datetime.now()
        elapsed = end_time - start_time
        print(f"\n完了: {self.output_file}")
        print(f"総処理時間: {elapsed}")


def main():
    # Windowsでの並列処理の問題を回避するための設定
    if sys.platform == 'win32':
        multiprocessing.freeze_support()
    
    parser = argparse.ArgumentParser(description='ファイル情報集約ツール')
    parser.add_argument('input_source', help='集約するルートディレクトリまたは開始URL')
    parser.add_argument('output_file', help='出力ファイルのパス')
    
    args = parser.parse_args()
    
    aggregator = FileAggregator(args.input_source, args.output_file)
    aggregator.aggregate()


if __name__ == "__main__":
    main()